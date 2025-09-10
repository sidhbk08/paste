import os
import fitz  # PyMuPDF
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import re
import xlsxwriter

# ----------------------
# Extract Text Functions
# ----------------------

def extract_pdf_text(pdf_path):
    """Extract text lines from PDF pages."""
    doc = fitz.open(pdf_path)
    lines = []
    for page in doc:
        text = page.get_text()
        lines.extend(text.split('\n'))
    doc.close()
    return lines

def extract_docx_text_with_headers_footers_tables(docx_path):
    """Extract text lines from DOCX headers, body paragraphs, tables, and footers."""
    doc = Document(docx_path)
    lines = []

    # Extract headers
    for rel in doc.part.rels.values():
        if rel.reltype == RT.HEADER:
            header = rel.target_part
            for p in header._element.xpath(".//w:p"):
                line = "".join(node.text for node in p.xpath(".//w:t") if node.text)
                if line.strip():
                    lines.append(line)

    # Extract body paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)

    # Extract tables (each row joined with ' | ')
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)

    # Extract footers
    for rel in doc.part.rels.values():
        if rel.reltype == RT.FOOTER:
            footer = rel.target_part
            for p in footer._element.xpath(".//w:p"):
                line = "".join(node.text for node in p.xpath(".//w:t") if node.text)
                if line.strip():
                    lines.append(line)

    return lines

# ----------------------
# HTML Report Generator
# ----------------------

html_template = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>PDF vs DOCX Comparison - {filename}</title>
    <style>
        body {{ font-family: Arial, sans-serif; display: flex; margin:0; padding:0; }}
        .pane {{ width: 50%; padding: 20px; overflow-y: scroll; height: 90vh; border-right: 1px solid #ccc; box-sizing: border-box; }}
        .highlight-missing {{ background-color: #fdd; }}
        .highlight-extra {{ background-color: #dfd; }}
        h2 {{ text-align: center; margin-top: 0; }}
        .paragraph {{ margin-bottom: 1em; white-space: pre-wrap; }}
    </style>
    <script>
        function syncScroll(el1, el2) {{
            el1.onscroll = function() {{ el2.scrollTop = el1.scrollTop; }};
            el2.onscroll = function() {{ el1.scrollTop = el2.scrollTop; }};
        }}
        window.onload = function() {{
            var left = document.getElementById("pdfPane");
            var right = document.getElementById("docxPane");
            syncScroll(left, right);
        }};
    </script>
</head>
<body>
    <div id="pdfPane" class="pane">
        <h2>PDF</h2>
        {pdf_content}
    </div>
    <div id="docxPane" class="pane">
        <h2>DOCX</h2>
        {docx_content}
    </div>
</body>
</html>
"""

def generate_html_report(filename, pdf_lines, docx_lines, missing_in_docx, missing_in_pdf, output_dir):
    """Generate a side-by-side HTML comparison with highlights and synchronized scrolling."""
    pdf_html = ""
    docx_html = ""

    def highlight_line(line, missing_set, css_class):
        words = re.findall(r'\w+', line.lower())
        if any(word in missing_set for word in words):
            return f'<div class="paragraph {css_class}">{line}</div>'
        return f'<div class="paragraph">{line}</div>'

    for line in pdf_lines:
        pdf_html += highlight_line(line, missing_in_docx, "highlight-missing") + "\n"

    for line in docx_lines:
        docx_html += highlight_line(line, missing_in_pdf, "highlight-extra") + "\n"

    html = html_template.format(
        filename=filename,
        pdf_content=pdf_html,
        docx_content=docx_html
    )

    html_path = os.path.join(output_dir, f"{filename}_comparison.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)
    return html_path

# ----------------------
# Comparison Logic
# ----------------------

def compare_word_sets(pdf_lines, docx_lines):
    """Return sets of words missing in DOCX and missing in PDF."""
    pdf_words = [w.lower() for line in pdf_lines for w in re.findall(r'\w+', line)]
    docx_words = [w.lower() for line in docx_lines for w in re.findall(r'\w+', line)]
    return set(pdf_words) - set(docx_words), set(docx_words) - set(pdf_words)

def compare_folder(folder_path):
    files = os.listdir(folder_path)
    pdf_files = [f for f in files if f.lower().endswith(".pdf")]

    # Create Excel report with two sheets
    workbook = xlsxwriter.Workbook(os.path.join(folder_path, "comparison_report.xlsx"))
    ws_docx = workbook.add_worksheet("Missing in DOCX")
    ws_pdf = workbook.add_worksheet("Missing in PDF")
    ws_docx.write_row(0, 0, ["Filename", "Missing Words"])
    ws_pdf.write_row(0, 0, ["Filename", "Missing Words"])

    row_docx = row_pdf = 1

    for pdf_file in pdf_files:
        base = pdf_file[:-4]
        docx_file = f"{base}.docx"
        pdf_path = os.path.join(folder_path, pdf_file)
        docx_path = os.path.join(folder_path, docx_file)

        if not os.path.exists(docx_path):
            print(f"⚠️ DOCX not found for {pdf_file}")
            continue

        print(f"\n🔍 Comparing {pdf_file} vs {docx_file}")
        pdf_lines = extract_pdf_text(pdf_path)
        docx_lines = extract_docx_text_with_headers_footers_tables(docx_path)

        missing_in_docx, missing_in_pdf = compare_word_sets(pdf_lines, docx_lines)

        if missing_in_docx:
            ws_docx.write(row_docx, 0, base)
            ws_docx.write(row_docx, 1, ", ".join(sorted(missing_in_docx)))
            row_docx += 1

        if missing_in_pdf:
            ws_pdf.write(row_pdf, 0, base)
            ws_pdf.write(row_pdf, 1, ", ".join(sorted(missing_in_pdf)))
            row_pdf += 1

        html_path = generate_html_report(base, pdf_lines, docx_lines, missing_in_docx, missing_in_pdf, folder_path)
        print(f"📄 HTML report saved to: {html_path}")

    workbook.close()
    print("\n✅ Excel report saved.")

# ----------------------
# Run the Tool
# ----------------------

if __name__ == "__main__":
    folder = input("📁 Enter the folder path containing PDF and DOCX files: ").strip()
    if os.path.isdir(folder):
        compare_folder(folder)
    else:
        print("❌ Invalid folder path.")

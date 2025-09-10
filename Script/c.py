from docx import Document

def copy_and_replace_text(source_path, template_path, start_text, end_text):
    # Load the source and template documents
    source_doc = Document(source_path)
    template_doc = Document(template_path)

    # Step 1: Extract the text and formatting from the source document
    copying = False
    copied_runs = []

    for paragraph in source_doc.paragraphs:
        if start_text in paragraph.text:
            copying = True  # Start copying text
        if copying:
            # Store the runs (formatted text) from the paragraph
            for run in paragraph.runs:
                copied_runs.append(run)
            if end_text in paragraph.text:
                break  # Stop copying after reaching the end text

    # Step 2: Find the start and end text in the template document
    start_found = False
    end_found = False
    for paragraph in template_doc.paragraphs:
        if start_text in paragraph.text:
            start_found = True
            # Clear the paragraph to replace it
            paragraph.clear()
            # Add the copied runs to the paragraph
            for run in copied_runs:
                new_run = paragraph.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.size = run.font.size
            # Add a new line after the copied text
            paragraph.add_run('\n')
        elif end_text in paragraph.text:
            end_found = True
            # Clear the end paragraph
            paragraph.clear()
        if start_found and end_found:
            break  # Stop after replacing the start and end text

    # Step 3: Save the modified template document with the same name
    template_doc.save(template_path)
    print(f"Updated document saved as: {template_path}")

# Get file paths from user input
source_path = input("Enter the path for the source document (TestS.docx): ")
template_path = input("Enter the path for the template document (TestT.docx): ")

# Define the start and end text
start_text = "You Have the Right to Appeal Our Decision"
end_text = "Plan share may include a reduction for sequestration"

# Call the function to copy and replace text
copy_and_replace_text(source_path, template_path, start_text, end_text)